"""پردازش ویس‌های دریافتی گروه‌های روبیکا — فاز ۶ سند.

اجرا: Celery task هر دقیقه یک‌بار (process-rubika-voices-every-minute در tasks.py).

جریان:
۱. پیام‌های voice با transcription=None از RubikaGroupMessage بخوان.
۲. فایل صوتی را با Whisper (local اگر موجود، وگرنه OpenAI API) به متن تبدیل کن.
۳. نتیجه را در RubikaGroupMessage.transcription ذخیره کن.
۴. به فایل Word روزانه اضافه کن:
       storage/transcriptions/rubika/YYYY-MM-DD.docx
   فرمت هر entry (مطابق سند):
       ─────────────────────────
       گروه: [نام]  |  فرستنده: [نام] — [شماره]  |  زمان: [ساعت]
       [متن]

سیاست whisper:
- اگر WHISPER_LOCAL_MODEL تنظیم شده: whisper محلی (بدون هزینه).
- در غیر این صورت: openai.Audio.transcriptions (API).
- هر دو مسیر به‌صورت try/except fallback هستند — اگر هر دو fail شد،
  transcription = "[پردازش ناموفق]" و پیام جهت retry باقی می‌ماند.

ملاحظه: whisper محلی CPU-intensive است. برای production با GPU/inference server
  بهتر عمل می‌کند. تنظیمات: WHISPER_LOCAL_MODEL (base|small|medium|large).
"""

from __future__ import annotations

import logging
import os
from datetime import date, datetime
from pathlib import Path

from sqlalchemy.orm import Session

from core_engine.models import RubikaGroupMessage

logger = logging.getLogger("core_engine.services.rubika_voice_processor")

PROJECT_ROOT = Path(__file__).resolve().parents[2]
TRANSCRIPTIONS_DIR = PROJECT_ROOT / "storage" / "transcriptions" / "rubika"

_FAILED_PLACEHOLDER = "[پردازش ناموفق]"
_BATCH_SIZE = 20  # حداکثر پیام در هر اجرای task


# ─── Whisper helpers ───────────────────────────────────────────────────────────


def _local_whisper_model():
    """مدل محلی Whisper را بارگذاری می‌کند (lazy, singleton)."""
    model_name = os.getenv("WHISPER_LOCAL_MODEL", "").strip()
    if not model_name:
        return None
    try:
        import whisper  # type: ignore[import-not-found]

        return whisper.load_model(model_name)
    except Exception:
        logger.exception("rubika_voice: بارگذاری مدل whisper محلی ناموفق بود — به API fallback می‌کند")
        return None


_WHISPER_MODEL = None


def _get_whisper_model():
    global _WHISPER_MODEL
    if _WHISPER_MODEL is None:
        _WHISPER_MODEL = _local_whisper_model()
    return _WHISPER_MODEL


async def _transcribe_file(file_path: str) -> str:
    """یک فایل صوتی را transcribe کرده و متن فارسی/فارسی+لاتین را برمی‌گرداند."""
    # تلاش ۱: whisper محلی
    model = _get_whisper_model()
    if model is not None:
        try:
            result = model.transcribe(file_path, language="fa")
            text = str(result.get("text", "")).strip()
            if text:
                logger.info("rubika_voice: transcribed locally (%d chars)", len(text))
                return text
        except Exception:
            logger.exception("rubika_voice: whisper local transcription failed for %s", file_path)

    # تلاش ۲: OpenAI Whisper API
    api_key = os.getenv("OPENAI_API_KEY", "").strip()
    if not api_key:
        logger.warning("rubika_voice: OPENAI_API_KEY تنظیم نشده — transcription ناموفق")
        return _FAILED_PLACEHOLDER

    try:
        from openai import AsyncOpenAI

        client = AsyncOpenAI(api_key=api_key)
        with open(file_path, "rb") as audio_file:
            response = await client.audio.transcriptions.create(
                model="whisper-1",
                file=audio_file,
                language="fa",
            )
        text = str(response.text).strip()
        logger.info("rubika_voice: transcribed via API (%d chars)", len(text))
        return text or _FAILED_PLACEHOLDER
    except Exception:
        logger.exception("rubika_voice: OpenAI API transcription failed for %s", file_path)
        return _FAILED_PLACEHOLDER


# ─── Word daily file ─────────────────────────────────────────────────────────


def _get_daily_docx_path(target_date: date | None = None) -> Path:
    d = target_date or date.today()
    return TRANSCRIPTIONS_DIR / f"{d.isoformat()}.docx"


def _format_received_at(received_at: datetime | None) -> str:
    if received_at is None:
        return "نامشخص"
    return received_at.strftime("%H:%M:%S")


def _append_to_daily_docx(
    *,
    group_name: str | None,
    sender_name: str | None,
    sender_phone: str | None,
    received_at: datetime | None,
    transcription: str,
) -> None:
    """یک transcription را به فایل Word روزانه اضافه می‌کند.

    اگر فایل امروز وجود نداشت، می‌سازد. مطابق فرمت سند:
        ─────────────────────────
        گروه: [نام]  |  فرستنده: [نام] — [شماره]  |  زمان: [ساعت]
        [متن]
    """
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    TRANSCRIPTIONS_DIR.mkdir(parents=True, exist_ok=True)
    docx_path = _get_daily_docx_path()

    if docx_path.exists():
        doc = Document(str(docx_path))
    else:
        doc = Document()
        title = doc.add_heading("پیام‌های صوتی روبیکا — " + date.today().isoformat(), level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    separator = doc.add_paragraph("─" * 40)
    separator.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in separator.runs:
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    group_str = group_name or "نامشخص"
    sender_str = sender_name or "نامشخص"
    phone_str = f" — {sender_phone}" if sender_phone else ""
    time_str = _format_received_at(received_at)

    meta = doc.add_paragraph(
        f"گروه: {group_str}  |  فرستنده: {sender_str}{phone_str}  |  زمان: {time_str}"
    )
    meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in meta.runs:
        run.font.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x33, 0x55, 0x88)

    body = doc.add_paragraph(transcription)
    body.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in body.runs:
        run.font.size = Pt(10)

    doc.save(str(docx_path))
    logger.info("rubika_voice: appended to %s", docx_path)


# ─── main service function ────────────────────────────────────────────────────


async def process_pending_voices(db: Session) -> dict:
    """Celery task از این تابع مستقیم فراخوانی می‌کند.

    Returns:
        dict با کلیدهای processed, failed, skipped_missing_file
    """
    rows = (
        db.query(RubikaGroupMessage)
        .filter(
            RubikaGroupMessage.message_type == "voice",
            RubikaGroupMessage.voice_file_path.isnot(None),
            RubikaGroupMessage.transcription.is_(None),
        )
        .limit(_BATCH_SIZE)
        .all()
    )

    if not rows:
        return {"processed": 0, "failed": 0, "skipped_missing_file": 0}

    processed = failed = skipped = 0

    for row in rows:
        file_path = row.voice_file_path or ""
        if not Path(file_path).exists():
            logger.warning("rubika_voice: فایل %s وجود ندارد — skip", file_path)
            # می‌گذاریم خالی بماند تا task بعدی دوباره تلاش نکند — placeholder می‌گذاریم
            row.transcription = _FAILED_PLACEHOLDER
            db.flush()
            skipped += 1
            continue

        try:
            text = await _transcribe_file(file_path)
            row.transcription = text
            db.flush()

            _append_to_daily_docx(
                group_name=row.group_name,
                sender_name=row.sender_name,
                sender_phone=row.sender_phone,
                received_at=row.received_at,
                transcription=text,
            )
            processed += 1

        except Exception:
            logger.exception(
                "rubika_voice: پردازش پیام %s ناموفق", row.id
            )
            row.transcription = _FAILED_PLACEHOLDER
            db.flush()
            failed += 1

    db.commit()
    logger.info(
        "rubika_voice: processed=%d failed=%d skipped_missing=%d", processed, failed, skipped
    )
    return {"processed": processed, "failed": failed, "skipped_missing_file": skipped}
